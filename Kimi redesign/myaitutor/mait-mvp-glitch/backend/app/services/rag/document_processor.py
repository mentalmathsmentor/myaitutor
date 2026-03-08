"""
Document processor for NESA syllabus PDFs and DOCX files.
Implements syllabus-aware chunking that preserves topic code structure.
"""
import re
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass

import pypdf
from docx import Document as DocxDocument

from .config import TOPIC_CODE_PATTERN, CONTENT_CODE_PATTERN, MAX_CHUNK_SIZE


@dataclass
class SyllabusChunk:
    """Represents a chunk of syllabus content with metadata."""
    id: str
    text: str
    topic_code: str
    content_code: Optional[str]
    topic_name: str
    year: str
    parent_topic: str
    source: str

    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "text": self.text,
            "metadata": {
                "topic_code": self.topic_code,
                "content_code": self.content_code or "",
                "topic_name": self.topic_name,
                "year": self.year,
                "parent_topic": self.parent_topic,
                "source": self.source
            }
        }


class DocumentProcessor:
    """Processes PDF and DOCX syllabus documents into semantic chunks."""

    # Topic code to metadata mapping (NSW HSC Mathematics Advanced)
    TOPIC_METADATA = {
        "MA-F1": {"name": "Working with Functions", "year": "11", "parent": "Functions"},
        "MA-F2": {"name": "Graphing Techniques", "year": "12", "parent": "Functions"},
        "MA-T1": {"name": "Trigonometry and Measure of Angles", "year": "11", "parent": "Trigonometric Functions"},
        "MA-T2": {"name": "Trigonometric Functions and Identities", "year": "11", "parent": "Trigonometric Functions"},
        "MA-T3": {"name": "Trigonometric Functions and Graphs", "year": "12", "parent": "Trigonometric Functions"},
        "MA-C1": {"name": "Introduction to Differentiation", "year": "11", "parent": "Calculus"},
        "MA-C2": {"name": "Differential Calculus", "year": "12", "parent": "Calculus"},
        "MA-C3": {"name": "Applications of Differentiation", "year": "12", "parent": "Calculus"},
        "MA-C4": {"name": "Integral Calculus", "year": "12", "parent": "Calculus"},
        "MA-E1": {"name": "Logarithms and Exponentials", "year": "11", "parent": "Exponential and Logarithmic Functions"},
        "MA-M1": {"name": "Modelling Financial Situations", "year": "12", "parent": "Financial Mathematics"},
        "MA-S1": {"name": "Probability and Discrete Probability Distributions", "year": "11", "parent": "Statistical Analysis"},
        "MA-S2": {"name": "Descriptive Statistics and Bivariate Data", "year": "12", "parent": "Statistical Analysis"},
        "MA-S3": {"name": "Random Variables", "year": "12", "parent": "Statistical Analysis"},
    }

    def __init__(self):
        self.chunks: List[SyllabusChunk] = []

    def process_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF file."""
        reader = pypdf.PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    def process_docx(self, docx_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text

    def extract_text(self, file_path: Path) -> str:
        """Extract text from PDF or DOCX file."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self.process_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self.process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _infer_topic_code(self, content_code: str) -> str:
        """Infer topic code from content code (e.g., C2.1 -> MA-C2)."""
        if len(content_code) >= 2:
            letter = content_code[0]
            number = content_code[1]
            return f"MA-{letter}{number}"
        return "MA-UNK"

    def chunk_by_content_code(self, text: str, source: str) -> List[SyllabusChunk]:
        """
        Chunk text by content codes (e.g., C2.1, F1.3).
        Falls back to topic codes if content codes not found.
        """
        chunks = []

        # Pattern to match content code headers like "C2.1" or "F1.3" followed by content
        # Match patterns like: "C2.1:", "C2.1 ", "F1.3:", etc.
        content_pattern = re.compile(r'([A-Z][0-9]\.[0-9]+):?\s*(.{0,100})', re.MULTILINE)

        # Find all content code sections
        matches = list(content_pattern.finditer(text))

        if not matches:
            # Fallback: try to chunk by topic codes (MA-X#)
            return self._chunk_by_topic_code(text, source)

        for i, match in enumerate(matches):
            content_code = match.group(1)  # e.g., "C2.1"
            section_title = match.group(2).strip()

            # Determine topic code from content code
            topic_code = self._infer_topic_code(content_code)

            # Get section text (from this match to next match or end)
            start_pos = match.start()
            end_pos = matches[i + 1].start() if i + 1 < len(matches) else min(start_pos + 3000, len(text))
            section_text = text[start_pos:end_pos].strip()

            # Limit chunk size
            if len(section_text) > MAX_CHUNK_SIZE * 4:  # Rough char limit
                section_text = section_text[:MAX_CHUNK_SIZE * 4] + "..."

            # Get metadata
            metadata = self.TOPIC_METADATA.get(topic_code, {
                "name": section_title or "Unknown Topic",
                "year": "Unknown",
                "parent": "Mathematics"
            })

            chunk = SyllabusChunk(
                id=f"{topic_code}.{content_code}",
                text=section_text,
                topic_code=topic_code,
                content_code=content_code,
                topic_name=metadata["name"],
                year=metadata["year"],
                parent_topic=metadata["parent"],
                source=source
            )
            chunks.append(chunk)

        return chunks

    def _chunk_by_topic_code(self, text: str, source: str) -> List[SyllabusChunk]:
        """Fallback chunking by topic codes (MA-X#)."""
        chunks = []

        # Pattern to match topic code headers
        topic_pattern = re.compile(r'(MA-[A-Z][0-9])\s*[:\-]?\s*(.{0,100})', re.MULTILINE)
        matches = list(topic_pattern.finditer(text))

        if not matches:
            # Last resort: create one big chunk
            chunk = SyllabusChunk(
                id=f"DOC-{source[:20]}",
                text=text[:MAX_CHUNK_SIZE * 4],
                topic_code="MA-GEN",
                content_code=None,
                topic_name="General Mathematics",
                year="Unknown",
                parent_topic="Mathematics",
                source=source
            )
            return [chunk]

        for i, match in enumerate(matches):
            topic_code = match.group(1)
            section_title = match.group(2).strip()

            start_pos = match.start()
            end_pos = matches[i + 1].start() if i + 1 < len(matches) else min(start_pos + 5000, len(text))
            section_text = text[start_pos:end_pos].strip()

            if len(section_text) > MAX_CHUNK_SIZE * 4:
                section_text = section_text[:MAX_CHUNK_SIZE * 4] + "..."

            metadata = self.TOPIC_METADATA.get(topic_code, {
                "name": section_title or "Unknown",
                "year": "Unknown",
                "parent": "Mathematics"
            })

            chunk = SyllabusChunk(
                id=topic_code,
                text=section_text,
                topic_code=topic_code,
                content_code=None,
                topic_name=metadata["name"],
                year=metadata["year"],
                parent_topic=metadata["parent"],
                source=source
            )
            chunks.append(chunk)

        return chunks

    def process_all(self, file_paths: List[Path]) -> List[SyllabusChunk]:
        """Process all syllabus files and return chunks."""
        all_chunks = []

        for file_path in file_paths:
            if not file_path.exists():
                print(f"Warning: File not found: {file_path}")
                continue

            try:
                print(f"Processing: {file_path.name}")
                text = self.extract_text(file_path)
                chunks = self.chunk_by_content_code(text, file_path.name)
                all_chunks.extend(chunks)
                print(f"  -> Created {len(chunks)} chunks")
            except Exception as e:
                print(f"Error processing {file_path}: {e}")

        # Deduplicate by ID (prefer longer content)
        seen = {}
        for chunk in all_chunks:
            if chunk.id not in seen or len(chunk.text) > len(seen[chunk.id].text):
                seen[chunk.id] = chunk

        final_chunks = list(seen.values())
        print(f"\nTotal unique chunks: {len(final_chunks)}")
        return final_chunks
